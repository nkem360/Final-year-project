"""
Knowledge base ingestion.
Run this script once (or via API endpoint) to build the FAISS vector store
from veterinary documents in knowledge_base/documents/.

Supported formats: .txt, .pdf, .docx, .md
"""

import logging
from pathlib import Path
from typing import List

from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

from core.settings import KNOWLEDGE_BASE_PATH
from ai.embeddings import create_vector_store

logger = logging.getLogger(__name__)


def load_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def load_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except ImportError:
        logger.warning("pypdf not installed — skipping PDF ingestion")
        return ""


def load_docx(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        return "\n".join(p.text for p in doc.paragraphs)
    except ImportError:
        logger.warning("python-docx not installed — skipping DOCX ingestion")
        return ""


LOADERS = {
    ".txt": load_text_file,
    ".md": load_text_file,
    ".pdf": load_pdf,
    ".docx": load_docx,
}


def load_documents() -> List[Document]:
    """Load all documents from the knowledge base directory."""
    base_path = Path(KNOWLEDGE_BASE_PATH)
    if not base_path.exists():
        logger.warning(f"Knowledge base directory not found: {KNOWLEDGE_BASE_PATH}")
        return []

    documents = []
    for path in base_path.rglob("*"):
        if path.is_file() and path.suffix.lower() in LOADERS:
            loader = LOADERS[path.suffix.lower()]
            content = loader(path)
            if content.strip():
                documents.append(Document(
                    page_content=content,
                    metadata={"source": path.name, "path": str(path)},
                ))
                logger.info(f"Loaded: {path.name} ({len(content)} chars)")

    logger.info(f"Loaded {len(documents)} documents from knowledge base")
    return documents


def split_documents(documents: List[Document]) -> List[Document]:
    """Split documents into chunks for embedding."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=512,
        chunk_overlap=50,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_documents(documents)
    logger.info(f"Split into {len(chunks)} chunks")
    return chunks


async def ingest_knowledge_base() -> int:
    """
    Full ingestion pipeline.
    Returns the number of chunks indexed.
    """
    documents = load_documents()
    if not documents:
        logger.warning("No documents found to ingest")
        return 0

    chunks = split_documents(documents)
    await create_vector_store(chunks)
    logger.info(f"Knowledge base ingestion complete: {len(chunks)} chunks indexed")
    return len(chunks)


if __name__ == "__main__":
    import asyncio
    logging.basicConfig(level=logging.INFO)
    count = asyncio.run(ingest_knowledge_base())
    print(f"Ingested {count} chunks into the vector store.")
